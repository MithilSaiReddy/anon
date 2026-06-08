import os
import sys
import tempfile
from pathlib import Path

import pytest

PROJECT_ROOT = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, str(PROJECT_ROOT))

os.environ.setdefault("HF_HUB_CACHE", str(PROJECT_ROOT / "models"))


@pytest.fixture
def temp_dir():
    with tempfile.TemporaryDirectory(prefix="anon_test_") as td:
        yield Path(td)


@pytest.fixture
def sample_text():
    return """
    Meeting with John Doe from Acme Corp in New York.
    His email is john@acme.com and phone is 555-0100.
    The project budget is $150000 and timeline is 12 months.
    """


@pytest.fixture
def sample_entity_map():
    return {
        "PERSON_1": "John Doe",
        "ORG_1": "Acme Corp",
        "LOC_1": "New York",
    }


@pytest.fixture
def sample_bridge_key():
    return {
        "app": "Anon Bridge",
        "version": "2.0",
        "multiplier": 3.0,
        "created_at": "2026-06-01T00:00:00",
        "original_filename": "test.pdf",
        "entities": {
            "PERSON_1": "John Doe",
            "ORG_1": "Acme Corp",
            "LOC_1": "New York",
        },
        "number_mappings": {
            "150000": 450000.0,
        },
    }


@pytest.fixture
def sample_docx_bytes():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Meeting with John Doe from Acme Corp in New York.")
    doc.add_paragraph("Budget: $150000 and timeline 12 months.")
    buf = tempfile.SpooledTemporaryFile()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
